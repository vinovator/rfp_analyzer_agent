import streamlit as st
from extractor.rfp_extractor import extract_rfp_info
from utils.file_loader import load_pdf, load_docx
import json
from io import BytesIO
from docx import Document

st.set_page_config(page_title="RFP Analyzer", page_icon=":mag_right:", layout="wide")


def render_value(value):
    """
    Render the value in a Streamlit-friendly format. (formatted dict, list or string)
    Return the formatted content as plain text string
    """
    content_lines = []
    if isinstance(value, dict):
        # If the value is a dictionary, display its contents
        for sub_key, sub_value in value.items():
            st.markdown("- **{}**: {}".format(sub_key.replace("_", " ").title(), sub_value))
            content_lines.append("- **{}**: {}".format(sub_key.replace("_", " ").title(), sub_value))
    elif isinstance(value, list):
        # If the value is a list, display each item
        for item in value:
            st.markdown("- {}".format(item))
            content_lines.append("- {}".format(item))
    else:
        # Otherwise, display the value directly
        st.markdown("- {}".format(value))
        content_lines.append("- {}".format(value))

    return "\n".join(content_lines)


# ==== UI Layout ====
st.title("RFP Analyzer AI Agent :mag_right:")
st.markdown("Analyze RFPs from **PDF** or **DOCX** or **text** and extract structured information.")

col1, col2 = st.columns(2)
with col1:
    uploaded_file = st.file_uploader("Upload RFP Document (PDF or DOCX)", type=["pdf", "docx"])
with col2:
    rfp_text = st.text_area("Or paste RFP text here", height=300)

text_to_process = ""

# ==== File Upload Handling ====
if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1].lower()
    if file_type == "pdf":
        # Load PDF file
        with st.spinner("Loading PDF..."):
            text_to_process = load_pdf(uploaded_file)
    elif file_type == "docx":
        # Load DOCX file
        with st.spinner("Loading DOCX..."):
            text_to_process = load_docx(uploaded_file)
    else:
        st.error("Unsupported file type. Please upload a PDF or DOCX file.")
elif rfp_text.strip():
    # Use text from text area if no file is uploaded
    text_to_process = rfp_text.strip()


# ==== RFP Analysis ====

# Submit button
if st.button("Analyze RFP"):
    if not text_to_process:
        st.warning("Please upload a fle or paste some RFP text before clicking Analyze RFP button.")
    else:
        with st.spinner("Extracting insights using GPT..."):
            raw_output = extract_rfp_info(text_to_process)
            
            if raw_output:
                # Attempt to parse the output as JSON
                raw_output = raw_output.strip("```").strip("json") # Clean up the output as json string is appended with ```json
                raw_output = raw_output.strip()  # Remove any leading/trailing whitespace
                rfp_data = json.loads(raw_output)
                st.success("RFP information extracted successfully!")

                export_text = [] # store formatted content for export
                
                for key, value in rfp_data.items():
                    st.subheader("{}".format(key.replace("_", " ").title()))
                    formaatted_content = render_value(value)
                    export_text.append("### {}\n{}".format(key.replace("_", " ").title(), formaatted_content))
                
                final_text = "\n".join(export_text)    

                # ==== Export Options ====
                st.markdown("### Export Options")
                st.subheader("Download Extracted summary")

                # JSON Export
                st.download_button(
                    label="Download as JSON",
                    data=json.dumps(rfp_data, indent=2),
                    file_name="rfp_data.json",
                    mime="application/json"
                )

                # Text Export
                txt_content = "\n".join([f"{key}: {value}" for key, value in rfp_data.items()])
                st.download_button(
                    label="Download as Text",
                    data=final_text,
                    file_name="rfp_data.txt",
                    mime="text/plain"
                )

                # DOCX Export
                doc = Document()
                doc.add_heading('RFP Summary', level=1)
                for section in export_text:
                    lines = section.strip().split('\n', 1)
                    doc.add_heading(lines[0].strip(), level=2)
                    if len(lines) > 1:
                        for line in lines[1].split('\n'):
                            doc.add_paragraph(line.strip())
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)  # Reset the stream position to the beginning
                st.download_button(
                    label="Download as DOCX",
                    data=doc_io,
                    file_name="rfp_data.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
      
            else:
                st.error("Failed to parse the RFP data. showing the raw output instead.")
                st.write(raw_output)
